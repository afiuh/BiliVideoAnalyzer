import os
import re
import requests
import time
import subprocess
from openpyxl import load_workbook
from openpyxl.styles import Font
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from config import DEEPSEEK_API_KEY, DEEPSEEK_API_URL

EXCEL_FILE = "./data/video_scores.xlsx"
WORD_DIR = "./data/word_reviews"
SUBTITLE_DIR = "./data/subtitles"

def call_deepseek(prompt, max_retries=2):
    """调用 DeepSeek API，返回回复文本"""
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 4000
    }
    for attempt in range(max_retries):
        try:
            resp = requests.post(DEEPSEEK_API_URL, json=payload, headers=headers, timeout=60)
            resp.raise_for_status()
            result = resp.json()
            return result['choices'][0]['message']['content']
        except Exception as e:
            print(f"  API 调用失败 (尝试 {attempt+1}/{max_retries}): {e}")
            time.sleep(2)
    return None

def clean_text(text):
    """去除所有 # 和 * 字符"""
    return re.sub(r'[#*]', '', text)

def generate_word(bvid, ai_review):
    os.makedirs(WORD_DIR, exist_ok=True)
    filepath = os.path.join(WORD_DIR, f"{bvid}.docx")
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.add_paragraph("AI评价：")
    for line in clean_text(ai_review).split('\n'):
        doc.add_paragraph(line)
    doc.save(filepath)
    return os.path.relpath(filepath, start="./data").replace("\\", "/")

def build_s_prompt(title, subtitle):
    """构建 S 档提示词"""
    total_chars = len(subtitle)
    target_len = max(total_chars // 2, 500)
    prompt = f"""你是一个专业的视频内容分析师，熟悉视频质量的分级标准。请对以下B站视频的文案进行分析，并回答七个问题。

【S档视频的核心特征】
- 内容极深入，具有系统性：视频围绕一个核心主题展开完整的知识框架，不是零散观点堆砌。
- 多维度论证，旁征博引：大量引用历史案例、经典著作、真实人物故事，论证层层递进。
- 原创性解读，有独特视角：提供超越常识的独特分析视角，给人启发感。
- 逻辑严密，结构清晰：有明确的叙事主线，分章节展开，前后呼应。
- 语言精炼，信息密度高：每句话都承载信息，无废话，篇幅长但不觉冗长。
- 有思想深度，引人思考：传递价值观或思维方式，能引发读者反思。
- 结合现实，有实用价值：联系现代人的职场、人生选择，让人感觉“有用”。

视频标题：{title}
文案全文：
{subtitle}

请基于以上信息，用流畅的中文输出你的评价，并**按顺序**回答以下问题：

1. **核心内容与观点**：这个视频主要讲了什么？它的核心观点或主题是什么？（请用1-2句话概括）
2. **观点价值**：这个观点有没有价值？为什么？（例如，是否具有启发性、实用性、独特性，或者提供了新的视角？）
3. **论证逻辑**：视频是如何论证自己的观点的？逻辑性强不强？（请分析其论证结构、证据使用、推理过程等，可指出优点或不足）
4. **内容系统性**：视频是否有完整的知识框架，还是零散观点？请举例说明。
5. **论证丰富性**：是否旁征博引？引用的案例和素材是否有力支撑了观点？
6. **思想深度**：是否传递了值得思考的价值观或思维方式？
7. **现实价值**：对现代人是否有实用启发？

**最终评价**：请用一段话总结你对这个视频的整体看法，并说明它是否符合S档的标准。

请确保你的评价字数不少于 {target_len} 字（原视频字数为 {total_chars}）。

请在最末尾单独一行以“是否符合S档：是”或“是否符合S档：否”的格式给出结论。"""
    return prompt

def build_a_prompt(title, subtitle):
    """构建 A 档提示词"""
    prompt = f"""你是一个专业的视频内容分析师，熟悉视频质量的分级标准。请对以下B站视频的文案进行分析，并回答六个问题。

【A档视频的核心特征】
- 主题明确，但未形成完整的知识体系：视频有清晰的主题，但内容以个人体验、故事叙述或单一角度分析为主，而非系统性的知识框架。
- **思辨性与逻辑性要求**：必须根据视频题材区分要求：
  - 若为**个人题材**（如情感、个人成长、生活感悟）：可放宽逻辑要求，允许情感主导，但需情感真挚、有共鸣感。
  - 若为**社会集体题材**（如历史解读、社会现象分析、行业揭秘、群体观察）：必须有清晰的逻辑论证（因果链条、证据支撑、结构严谨），不能仅凭个人情感体验。
- **格局要求**：无论何种题材，均不得有“小资产阶级小家子气”——即无病呻吟、矫揉造作、格局狭小、过度关注个人琐碎情绪而无社会关怀或普世价值的内容。此类视频即使符合其他特征，也应降档。
- 情感或个人体验色彩（适用于个人题材）：能够引发情感共鸣，或带有强烈的个人视角。
- 故事性或叙事性强（适用于个人题材）：通过讲述故事来传递观点，有完整的情节线。
- 分析角度单一但深入（适用于社会题材）：聚焦于某个具体问题或人物，进行深入剖析，但缺乏多维度旁征博引。
- 信息密度适中，逻辑清晰：内容充实但不密集，逻辑结构清楚。
- 有思想启发但未成体系：能提供有价值的视角或感悟，但不足以作为系统性学习材料。

【与S档的主要区别】
- S档具有完整的知识框架和系统性，A档缺乏系统性。
- S档旁征博引多案例，A档案例相对集中或单一。
- S档信息密度极高，A档信息密度适中。
- S档更偏理性启发，A档更偏情感共鸣或个人感悟。

【与B档（资讯汇编）的主要区别】
- B档可能信息密度高、逻辑清晰，但多为事实陈述、行业揭秘、知识科普，缺乏原创性深度见解或独特视角。
- A档必须包含超越资讯整合的、具有个人思考深度的观点。

视频标题：{title}
文案全文：
{subtitle}

请基于以上信息，用流畅的中文输出你的评价，并**按顺序**回答以下问题：

1. **核心内容与观点**：这个视频主要讲了什么？它的核心观点或主题是什么？（请用1-2句话概括）
2. **观点价值**：这个观点有没有价值？为什么？（例如，是否具有启发性、实用性、独特性，或者提供了新的视角？）
3. **论证逻辑**：视频是如何论证自己的观点的？逻辑性强不强？（请分析其论证结构、证据使用、推理过程等，可指出优点或不足）
4. **原创性见解**：视频是否提供了超越常识的、具有个人深度的观点？还是仅仅是事实和资讯的罗列？
5. **情感/体验色彩**：视频是否具有明显的情感或个人体验色彩？如果有，请描述。
6. **思想启发**：视频是否提供了有价值的视角或感悟？

**最终评价**：请用一段话总结你对这个视频的整体看法，并说明它是否符合A档标准。如果符合，请进一步指出它是属于“A(体验)”还是“A(分析)”，并简述理由。

请确保你的评价字数不少于 1200 字。

请在最末尾单独一行以“是否符合A档：是”或“是否符合A档：否”的格式给出结论。如果符合，请在下一行以“具体档次：A(体验)”或“具体档次：A(分析)”的格式给出具体档次。"""
    return prompt
# ================== 新增：立场检测与批判函数 ==================
def build_stance_check_prompt(title, subtitle):
    """构建立场检测提示词"""
    prompt = f"""你是一个政治立场审查员。请分析以下视频文案，判断是否存在以下问题：
- 断章取义地歪曲事实
- 使用侮辱性词汇攻击国家、政府、军队或相关政策
- 凭空捏造不存在的负面事件或数据
- 阴阳怪气地讽刺或影射

如果存在上述任何一种情况，请在最后一行输出“立场判断：是”；否则输出“立场判断：否”。无需输出其他分析。

视频标题：{title}
文案：
{subtitle}
"""
    return prompt

def build_critique_prompt(title, subtitle):
    """构建暴力批判提示词"""
    prompt = f"""你是一个立场坚定的爱国者。以下视频文案存在恶意歪曲事实、侮辱国家/政府/军队、阴阳怪气等问题。请你以最严厉的语气驳斥其错误逻辑，揭露其荒谬之处，并指出正确的立场。请用批判性的语言，有理有据地反驳，但不要过度情绪化。

视频标题：{title}
文案：
{subtitle}

请直接开始你的批判："""
    return prompt

def force_close_office_apps():
    """强制关闭所有 Excel 和 Word 进程（Windows）"""
    try:
        # 关闭 Excel
        subprocess.run('taskkill /f /im excel.exe', shell=True, capture_output=True)
        # 关闭 Word
        subprocess.run('taskkill /f /im winword.exe', shell=True, capture_output=True)
        print("已尝试关闭 Excel 和 Word 进程。")
    except Exception as e:
        print(f"关闭进程时出错: {e}")

def main():
    force_close_office_apps()   # 在脚本一开始就尝试关闭 Office 进程
    # 检查 Excel 是否存在
    if not os.path.exists(EXCEL_FILE):
        print(f"错误：找不到 {EXCEL_FILE}，请先运行评分脚本。")
        return

    wb = load_workbook(EXCEL_FILE)
    ws = wb.active
    # 获取表头
    headers = [cell.value for cell in ws[1]]   # 定义 headers
    try:
        col_bvid = headers.index("BV号") + 1
    except ValueError:
        print("Excel 表头缺少 BV 号列，无法处理。")
        return
    try:
        col_title = headers.index("标题") + 1
    except ValueError:
        col_title = None  # 标题列不存在，后续将标题设为未知
    try:
        col_rating = headers.index("最终评级") + 1
    except ValueError:
        print("Excel 表头缺少最终评级列，无法处理。")
        return

    # 如果不存在“Word文件”列，则添加
    col_word = None
    if "Word文件" not in headers:
        col_word = len(headers) + 1
        ws.cell(row=1, column=col_word, value="Word文件")
    else:
        col_word = headers.index("Word文件") + 1

    # 遍历数据行（从第二行开始）
    for row in range(2, ws.max_row + 1):
        bvid = ws.cell(row=row, column=col_bvid).value
        if col_title is not None:
            title = ws.cell(row=row, column=col_title).value
        else:
            title = "未知标题"
        rating = ws.cell(row=row, column=col_rating).value

        # 只处理 S/A 档
        if not (rating and (rating.startswith('S') or rating.startswith('A'))):
            continue

        # 如果已生成 Word，跳过
        if ws.cell(row=row, column=col_word).value:
            print(f"跳过已处理：{bvid}")
            continue

        print(f"正在处理 S/A 档视频：{title}")
        # 读取字幕文件
        txt_path = os.path.join(SUBTITLE_DIR, f"{bvid}.txt")
        if not os.path.exists(txt_path):
            print(f"  字幕文件不存在，跳过")
            continue
        with open(txt_path, "r", encoding="utf-8") as f:
            subtitle = f.read()
            # ========== 新增：立场检测 ==========
            stance_prompt = build_stance_check_prompt(title, subtitle)
            stance_reply = call_deepseek(stance_prompt)
            if not stance_reply:
                print(f"  立场检测 API 调用失败，跳过视频 {bvid}")
                continue

            # 解析立场判断
            stance_passed = True  # 默认通过
            if "立场判断：是" in stance_reply:
                stance_passed = False
            elif "立场判断：否" in stance_reply:
                stance_passed = True
            else:
                # 未按格式输出，记录警告后视为通过（保守处理）
                print(f"  警告：立场检测回复格式异常，将按通过处理。回复片段：{stance_reply[:100]}")
                stance_passed = True

            if not stance_passed:
                # 立场不正 → 暴力批判模式
                print(f"  视频 {bvid} 立场不正，启动批判模式...")
                critique_prompt = build_critique_prompt(title, subtitle)
                critique_reply = call_deepseek(critique_prompt)
                if critique_reply:
                    word_path = generate_word(bvid,critique_reply)
                    # 更新 Excel：写入 Word 路径，并将评级改为 X
                    cell = ws.cell(row=row, column=col_word)
                    cell.value = "打开文档"
                    cell.hyperlink = os.path.abspath(os.path.join(".", "data", word_path))
                    cell.font = Font(color="0000FF", underline="single")
                    ws.cell(row=row, column=col_rating).value = "X"
                    wb.save(EXCEL_FILE)
                    print(f"  已生成批判文档，评级改为 X")
                else:
                    print(f"  批判生成失败，跳过")
                continue  # 不再进行正常审核
            # ========== 立场检测结束 ==========
        # 根据评级选择提示词
        if rating.startswith('S'):
            prompt = build_s_prompt(title, subtitle)
        else:
            prompt = build_a_prompt(title, subtitle)

        # 调用 API
        reply = call_deepseek(prompt)
        if not reply:
            print(f"  API 调用失败，跳过")
            continue

        # 生成 Word 文档
        word_path = generate_word(bvid, reply)
        cell = ws.cell(row=row, column=col_word)
        cell.value = "打开文档"
        cell.hyperlink = os.path.abspath(os.path.join(".", "data", word_path))
        cell.font = Font(color="0000FF", underline="single")

        # 立即保存 Excel，避免意外丢失
        wb.save(EXCEL_FILE)
        print(f"  已生成 Word：{word_path}")

        # 礼貌延时
        time.sleep(1)
    # 备份旧文件（移动原文件）
    EXCEL_BACKUP_DIR = "./data/excel_backups"
    os.makedirs(EXCEL_BACKUP_DIR, exist_ok=True)
    timestamp = time.strftime("%Y%m%d_%H%M")
    backup_path = os.path.join(EXCEL_BACKUP_DIR, f"{timestamp}.xlsx")
    os.replace(EXCEL_FILE, backup_path)
    print(f"旧文件已备份为: {backup_path}")
    print("AI 审核完成，Excel 已更新。")

if __name__ == "__main__":
    main()